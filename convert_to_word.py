import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_markdown(md_text):
    lines = md_text.split('\n')
    elements = []
    in_code_block = False
    code_content = []
    
    for line in lines:
        if line.strip().startswith('```'):
            if in_code_block:
                elements.append(('code', '\n'.join(code_content)))
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        if line.strip() == '':
            elements.append(('empty', ''))
            continue
        
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
        elif line.startswith('#### '):
            elements.append(('h4', line[5:].strip()))
        elif line.startswith('- '):
            elements.append(('list', line[2:].strip()))
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            elements.append(('ordered_list', text))
        elif line.startswith('> '):
            elements.append(('quote', line[2:].strip()))
        elif line.startswith('|') and '|' in line[1:]:
            elements.append(('table_row', line))
        elif line.startswith('---'):
            elements.append(('hr', ''))
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip().strip('*')
            elements.append(('bold', text))
        else:
            elements.append(('paragraph', line))
    
    return elements


def add_inline_formatting(paragraph, text):
    pattern = r'\*\*(.+?)\*\*'
    parts = re.split(pattern, text)
    
    for i, part in enumerate(parts):
        if i % 2 == 1:
            run = paragraph.add_run(part)
            run.bold = True
        else:
            pattern_code = r'`(.+?)`'
            code_parts = re.split(pattern_code, part)
            for j, code_part in enumerate(code_parts):
                if j % 2 == 1:
                    run = paragraph.add_run(code_part)
                    run.font.name = 'Consolas'
                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                else:
                    paragraph.add_run(code_part)


def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    elements = parse_markdown(md_text)
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    table_rows = []
    in_table = False
    
    for elem_type, content in elements:
        if elem_type == 'table_row':
            if '---' in content and all(c.strip() in ('---', '|') or c.strip().startswith(':') for c in content.split('|')[1:-1]):
                continue
            cells = [c.strip() for c in content.split('|')[1:-1]]
            table_rows.append(cells)
            in_table = True
            continue
        else:
            if in_table and table_rows:
                if table_rows:
                    num_cols = len(table_rows[0])
                    num_rows = len(table_rows)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for i, row in enumerate(table_rows):
                        for j, cell_text in enumerate(row):
                            cell = table.cell(i, j)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            add_inline_formatting(p, cell_text)
                            if i == 0:
                                for run in p.runs:
                                    run.bold = True
                
                table_rows = []
                in_table = False
        
        if elem_type == 'h1':
            p = doc.add_heading(level=1)
            run = p.add_run(content)
            run.font.size = Pt(22)
            run.font.bold = True
        elif elem_type == 'h2':
            p = doc.add_heading(level=2)
            run = p.add_run(content)
            run.font.size = Pt(18)
            run.font.bold = True
        elif elem_type == 'h3':
            p = doc.add_heading(level=3)
            run = p.add_run(content)
            run.font.size = Pt(15)
            run.font.bold = True
        elif elem_type == 'h4':
            p = doc.add_heading(level=4)
            run = p.add_run(content)
            run.font.size = Pt(13)
            run.font.bold = True
        elif elem_type == 'paragraph':
            if content.strip():
                p = doc.add_paragraph()
                add_inline_formatting(p, content)
        elif elem_type == 'list':
            p = doc.add_paragraph(style='List Bullet')
            add_inline_formatting(p, content)
        elif elem_type == 'ordered_list':
            p = doc.add_paragraph(style='List Number')
            add_inline_formatting(p, content)
        elif elem_type == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(content)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif elem_type == 'code':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Inches(0.2)
        elif elem_type == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
        elif elem_type == 'hr':
            p = doc.add_paragraph()
            p.add_run('_' * 50)
        elif elem_type == 'empty':
            pass
    
    if in_table and table_rows:
        num_cols = len(table_rows[0])
        num_rows = len(table_rows)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(table_rows):
            for j, cell_text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                add_inline_formatting(p, cell_text)
                if i == 0:
                    for run in p.runs:
                        run.bold = True
    
    doc.save(docx_path)
    print(f'转换完成: {docx_path}')


if __name__ == '__main__':
    md_to_docx('code_review.md', 'AI代码审查报告.docx')
    md_to_docx('summary.md', '个人实训总结报告.docx')
